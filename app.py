import streamlit as st
from docx import Document
from zipfile import ZipFile
import os
import tempfile
from io import BytesIO
from pdf2docx import Converter

# Function to combine Word documents with folder name as student name
def combine_word_documents(docs_with_names):
    combined_doc = Document()

    for name, doc_bytes in docs_with_names:
        # Truncate the name at the first underscore (if present)
        truncated_name = name.split('_')[0]

        # Add the student's name (truncated) at the start of each submission
        combined_doc.add_paragraph(f"STUDENT NAME: {truncated_name}")
        combined_doc.add_paragraph("")  # Add an empty line for spacing

        # Process the document
        sub_doc = Document(BytesIO(doc_bytes))
        for element in sub_doc.element.body:
            combined_doc.element.body.append(element)

        # Add a page break after each submission (optional)
        combined_doc.add_page_break()

    return combined_doc

# Function to convert PDF to Word document
def convert_pdf_to_word(pdf_bytes):
    with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as temp_pdf:
        temp_pdf.write(pdf_bytes)
        temp_pdf.flush()

        output = BytesIO()
        converter = Converter(temp_pdf.name)
        converter.convert(output)
        converter.close()

        output.seek(0)
        return output.getvalue()

# Function to process files from a ZIP and use folder name as submission name
def process_zip_file(uploaded_file):
    try:
        # Ensure the file is read as bytes
        zip_bytes = uploaded_file.read()
        with ZipFile(BytesIO(zip_bytes), 'r') as z:
            with tempfile.TemporaryDirectory() as tempdir:
                z.extractall(tempdir)
                processed_docs = []
                error_occurred = False

                # Loop through all extracted folders and files
                for root, dirs, files in os.walk(tempdir):
                    folder_name = os.path.basename(root)  # Get the folder name
                    for file in files:
                        try:
                            if file.endswith('.docx') or file.endswith('.pdf'):
                                file_path = os.path.join(root, file)
                                with open(file_path, 'rb') as f:
                                    if file.endswith('.pdf'):
                                        processed_docs.append((folder_name, convert_pdf_to_word(f.read())))  # Use folder name
                                    else:
                                        processed_docs.append((folder_name, f.read()))  # Use folder name
                        except Exception as e:
                            st.error(f"Error processing file {file} in folder {folder_name}: {str(e)}")
                            error_occurred = True
                return processed_docs, error_occurred

    except Exception as e:
        st.error(f"Error processing ZIP file: {str(e)}")
        return [], True

# Function to process direct file uploads
def process_files(files):
    processed_docs = []
    for file in files:
        if file.type == 'application/pdf':
            processed_docs.append(('Direct Upload', convert_pdf_to_word(file.getvalue())))
        else:
            processed_docs.append(('Direct Upload', file.getvalue()))
    return processed_docs

# Initialize session state variables
if 'combined_document' not in st.session_state:
    st.session_state['combined_document'] = None

# Streamlit UI
st.title('Word Document Combiner')
st.markdown("Word Document Combiner is an easy-to-use tool to merge multiple Word documents into one.")

upload_choice = st.radio("Choose your upload method", ('Zip File', 'Word Files'))

if upload_choice == 'Zip File':
    uploaded_file = st.file_uploader("Upload ZIP file", type=['zip'])
    if st.button('Combine Documents from ZIP') and uploaded_file:
        docs_with_names, error_occurred = process_zip_file(uploaded_file)
        if docs_with_names and not error_occurred:
            st.session_state['combined_document'] = combine_word_documents(docs_with_names)

elif upload_choice == 'Word Files':
    uploaded_files = st.file_uploader("Upload Word files", accept_multiple_files=True, type=['docx'])
    if st.button('Combine Word Documents') and uploaded_files:
        docs_with_names = process_files(uploaded_files)
        st.session_state['combined_document'] = combine_word_documents(docs_with_names)

# Export options
if st.session_state['combined_document']:
    export_format = st.selectbox("Select export format", ("Word", "PDF", "Text"))

    if st.button('Export Combined Document'):
        combined_document = st.session_state['combined_document']

        if export_format == "Word":
            file_stream = BytesIO()
            combined_document.save(file_stream)
            file_stream.seek(0)
            st.download_button(label="Download Combined Document",
                               data=file_stream,
                               file_name="combined_document.docx",
                               mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

        elif export_format == "PDF":
            # PDF conversion logic here
            pass

        elif export_format == "Text":
            text_stream = BytesIO()
            for paragraph in combined_document.paragraphs:
                text_stream.write(paragraph.text.encode('utf-8') + b'\n')
            text_stream.seek(0)
            st.download_button(label="Download Combined Text",
                               data=text_stream,
                               file_name="combined_document.txt",
                               mime="text/plain")


st.markdown("""

## **How to Use the App**

### **Getting Started**

*   You don’t need to install anything on your computer. The app runs in a web browser.
*   Make sure you have all your Word documents ready. You can have them as separate files or put together in a ZIP file.

### **Combining Your Documents**

1.  **Open the App**: Go to the web link where the Word Document Combiner is hosted.
2.  **Choose Your Upload Method**:
    *   If your documents are in a ZIP file, select the 'Zip File' option.
    *   If you have individual Word files, select the 'Word Files' option.
3.  **Upload Your Files**: Click on the upload area to select files from your computer, or drag and drop them into the box.
4.  **Combine Your Documents**:
    *   If you uploaded a ZIP file, click ‘Combine Documents from ZIP’.
    *   If you uploaded Word files, click ‘Combine Word Documents’.
5.  **Wait for the Process to Complete**: The app will merge your documents and let you know when it's done.

### **Downloading Your Combined Document**

*   Once your documents are combined, you can choose how you want to download the combined document. You can download it as a Word document, a PDF, or a text file.
*   Click on the download button for the format you want, and the file will be saved to your computer.

### **Troubleshooting**

*   If you get an error message or something doesn’t work, check to make sure that you’re uploading the right type of file (.docx or .zip).
*   If there are more than one Word document in a folder within your ZIP file, the app will only combine the first one it finds.

## **Need Help?**

If you have any questions or need help using the app, don’t hesitate to reach out for support.
""")
